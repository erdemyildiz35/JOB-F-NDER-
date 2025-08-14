import yaml
from docx import Document
from pathlib import Path

def generate_cv(lang="en", job_title="General Application"):
    """profile.yaml'dan oku ve CV üret."""
    with open("profile.yaml", "r", encoding="utf-8") as f:
        profile = yaml.safe_load(f)

    doc = Document()
    doc.add_heading(profile["name"], 0)
    doc.add_paragraph(f"Email: {profile['email']}")
    doc.add_paragraph(f"Phone: {profile['phone']}")
    doc.add_paragraph(f"LinkedIn: {profile['linkedin']}")
    doc.add_paragraph(f"GitHub: {profile['github']}")
    doc.add_paragraph(profile["address"])

    # Summary
    summary_key = "summary_en" if lang == "en" else "summary_de"
    doc.add_heading("Profile", level=1)
    doc.add_paragraph(profile[summary_key])

    # Skills
    skills_key = "skills_en" if lang == "en" else "skills_de"
    doc.add_heading("Skills", level=1)
    for skill in profile[skills_key]:
        doc.add_paragraph(f"- {skill}")

    # Experience
    exp_key = "experience_en" if lang == "en" else "experience_de"
    doc.add_heading("Experience", level=1)
    for exp in profile[exp_key]:
        doc.add_paragraph(f"{exp['role']} - {exp['company']} ({exp['period']})")
        for detail in exp["details"]:
            doc.add_paragraph(f"  • {detail}")

    # Education
    edu_key = "education_en" if lang == "en" else "education_de"
    doc.add_heading("Education", level=1)
    for edu in profile[edu_key]:
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['period']})")

    # Save file
    Path("data/cvs").mkdir(parents=True, exist_ok=True)
    file_path = f"data/cvs/CV_{lang.upper()}_{job_title.replace(' ', '_')}.docx"
    doc.save(file_path)
    print(f"CV saved: {file_path}")
